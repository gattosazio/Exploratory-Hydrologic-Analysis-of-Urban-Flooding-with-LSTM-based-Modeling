import docx

doc = docx.Document("yaaaaas.docx")

# Let's find index ranges
# Section 4.5.3 (Data Preprocessing) is around P518
# Section 5.9 (Hydrologic Analysis Overview) is around P893

def extract_range(start_idx, end_idx, out_file):
    lines = []
    for idx in range(start_idx, min(end_idx, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        text = p.text.strip()
        if text:
            lines.append(f"P{idx} [{p.style.name}]: {text}")
    with open(out_file, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Wrote P{start_idx} to P{end_idx} to {out_file}")

# Methodology for Obj 4
extract_range(510, 652, "methodology_obj4.txt")

# Results for Obj 4
extract_range(890, 1085, "results_obj4.txt")
