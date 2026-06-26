import docx

doc = docx.Document("yaaaaas.docx")
print("Total paragraphs:", len(doc.paragraphs))

# Let's search for some keywords like EDA, preprocessing, rainfall event, objective 4, hydrologic
keywords = ["eda", "exploratory data", "preprocess", "structuring", "rainfall event", "objective", "hydrologic"]

print("\n--- Search by Keywords ---")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    for kw in keywords:
        if kw in text.lower():
            print(f"P{idx} [{p.style.name}]: {text[:150]}...")
            break

print("\n--- Headings in yaaaaas.docx ---")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    if p.style.name.startswith("Heading"):
        print(f"P{idx} [{p.style.name}]: {text}")
